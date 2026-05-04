import json
import math
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent.parent
ASSETS = ROOT / ".report_assets"
DIAGRAMS = ASSETS / "diagrams"
SCREENSHOTS = ASSETS / "screenshots"
TMP = ASSETS / "tmp"
OUT = ROOT / "prediplomna_praktyka_hohosha_io.docx"

FONT = "/System/Library/Fonts/Supplemental/Times New Roman.ttf"
FONT_BOLD = "/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf"


def font(size: int, bold: bool = False):
    return ImageFont.truetype(FONT_BOLD if bold else FONT, size)


def set_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(14)
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(1.25)
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_page_number(footer)


def add_paragraph(doc, text="", bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True, size=14):
    p = doc.add_paragraph()
    p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    else:
        p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    return p


def add_heading(doc, text):
    p = add_paragraph(doc, text.upper(), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_table_caption(doc, text):
    add_paragraph(doc, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)


def add_figure_caption(doc, text):
    add_paragraph(doc, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.line_spacing = 1.5
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(f"• {item}")
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(14)


def add_image(doc, path: Path, width_cm=15.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))


def make_canvas(name, width=1800, height=1100):
    DIAGRAMS.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    return img, draw, DIAGRAMS / name


def box(draw, xy, text, fill="#f4f8fb", outline="#1f4e79", title=False):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=18, outline=outline, width=4, fill=fill)
    f = font(30 if title else 24, bold=title)
    margin = 20
    draw.multiline_text((x1 + margin, y1 + margin), text, fill="black", font=f, spacing=10)


def arrow(draw, start, end, text=None):
    draw.line([start, end], fill="#444", width=4)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    for offset in (-0.4, 0.4):
        x = end[0] - 20 * math.cos(angle + offset)
        y = end[1] - 20 * math.sin(angle + offset)
        draw.line([end, (x, y)], fill="#444", width=4)
    if text:
        mx = (start[0] + end[0]) / 2
        my = (start[1] + end[1]) / 2
        draw.text((mx, my - 25), text, fill="black", font=font(22))


def ellipse(draw, xy, text):
    draw.ellipse(xy, outline="#7a3e00", width=4, fill="#fff4e5")
    draw.multiline_text((xy[0] + 30, xy[1] + 35), text, fill="black", font=font(24), spacing=8)


def generate_use_case():
    img, draw, out = make_canvas("use_case.png", 1800, 1200)
    box(draw, (720, 80, 1170, 1080), "Інформаційна система\nспортивних секцій", fill="#f8fbff", title=True)
    draw.text((150, 150), "Користувач", fill="black", font=font(28, True))
    draw.text((1400, 120), "Адміністратор", fill="black", font=font(28, True))
    draw.text((1430, 760), "Тренер", fill="black", font=font(28, True))

    ellipse(draw, (790, 150, 1100, 250), "Реєстрація\nта авторизація")
    ellipse(draw, (760, 290, 1130, 390), "Перегляд секцій\nта тренувань")
    ellipse(draw, (760, 430, 1130, 530), "Запис / повторний\nзапис на тренування")
    ellipse(draw, (760, 570, 1130, 670), "Перегляд профілю\nта історії записів")
    ellipse(draw, (760, 710, 1130, 810), "Керування секціями,\nкористувачами і тренуваннями")
    ellipse(draw, (760, 850, 1130, 950), "Підтвердження, скасування\nта завершення тренувань")

    arrow(draw, (280, 190), (790, 200))
    arrow(draw, (280, 340), (760, 340))
    arrow(draw, (280, 480), (760, 480))
    arrow(draw, (280, 620), (760, 620))
    arrow(draw, (1510, 200), (1130, 760), "CRUD")
    arrow(draw, (1510, 280), (1130, 900), "статуси")
    arrow(draw, (1510, 830), (1130, 900))
    arrow(draw, (1510, 930), (1130, 480), "реєстрації")
    img.save(out)
    return out


def generate_class_diagram():
    img, draw, out = make_canvas("class_diagram.png", 1800, 1200)
    box(draw, (80, 80, 500, 390), "User\n\n+ id\n+ name\n+ email\n+ password\n+ role\n+ deleted_at\n\nregistrations()\ncoach()", title=True)
    box(draw, (660, 80, 1120, 330), "Coach\n\n+ id\n+ user_id\n+ bio\n+ phone\n+ specialization\n\nuser()\nsports()", title=True)
    box(draw, (1270, 80, 1720, 360), "Sport\n\n+ id\n+ name\n+ description\n+ location\n+ coach_id\n\ncoach()\ntrainings()\nregistrations()", title=True)
    box(draw, (220, 650, 700, 1040), "Registration\n\n+ id\n+ user_id\n+ training_id\n+ status\n+ created_at\n\nuser()\ntraining()\nSTATUS_PENDING\nSTATUS_APPROVED\nSTATUS_CANCELLED\nSTATUS_REJECTED\nSTATUS_ATTENDED\nSTATUS_NO_SHOW", title=True)
    box(draw, (1040, 650, 1560, 1040), "Training\n\n+ id\n+ sport_id\n+ date\n+ time\n+ place\n+ notes\n+ is_cancelled\n+ is_completed\n+ cancelled_at\n+ completed_at\n\nsport()\nregistrations()", title=True)

    arrow(draw, (500, 200), (660, 200), "1 : 1")
    arrow(draw, (1120, 200), (1270, 200), "1 : N")
    arrow(draw, (340, 390), (420, 650), "1 : N")
    arrow(draw, (1450, 360), (1320, 650), "1 : N")
    arrow(draw, (700, 840), (1040, 840), "N : 1")
    img.save(out)
    return out


def generate_component_diagram():
    img, draw, out = make_canvas("component_diagram.png", 1800, 1100)
    box(draw, (80, 120, 480, 320), "Клієнтський браузер\n\nVue 3 + Inertia.js\nЛокалізація\nКалендар тренувань", title=True)
    box(draw, (620, 70, 1180, 380), "Web-рівень Laravel\n\nМаршрути web.php\nMiddleware: auth, admin, coach, locale\nКонтролери Auth, Sport, Training,\nRegistration, Admin, Coach", title=True)
    box(draw, (620, 470, 1180, 850), "Бізнес-логіка та моделі\n\nUser, Coach, Sport, Training, Registration\nEloquent ORM\nВалідація запитів\nАгрегації статистики та статусів", title=True)
    box(draw, (1320, 150, 1710, 360), "MySQL 8\n\nтаблиці користувачів,\nсекцій, тренувань,\nреєстрацій, тренерів", title=True)
    box(draw, (1320, 500, 1710, 720), "Redis\n\nкеш, сесії,\nчерги", title=True)
    box(draw, (80, 520, 480, 760), "Docker-оточення\n\napp, nginx, db,\nredis, vite, queue", title=True)

    arrow(draw, (480, 220), (620, 220), "HTTP / Inertia")
    arrow(draw, (900, 380), (900, 470))
    arrow(draw, (1180, 250), (1320, 250), "SQL")
    arrow(draw, (1180, 660), (1320, 620), "cache / queue")
    arrow(draw, (480, 640), (620, 640), "контейнери")
    arrow(draw, (280, 520), (280, 320))
    img.save(out)
    return out


def generate_sequence_diagram():
    img, draw, out = make_canvas("sequence_diagram.png", 1800, 1200)
    actors = [
        ("Користувач", 160),
        ("Vue/Inertia", 520),
        ("RegistrationController", 930),
        ("Registration", 1320),
        ("MySQL", 1640),
    ]
    for name, x in actors:
        draw.text((x - 70, 40), name, fill="black", font=font(24, True))
        draw.line([(x, 90), (x, 1080)], fill="#777", width=2)

    steps = [
        (160, 520, 140, "Натискає «Записатися»"),
        (520, 930, 240, "POST /trainings/{id}/register"),
        (930, 1320, 340, "Перевірка статусу\nта наявності запису"),
        (1320, 1640, 470, "SELECT / INSERT / UPDATE"),
        (1640, 1320, 590, "Підтвердження збереження"),
        (1320, 930, 700, "Результат моделі"),
        (930, 520, 820, "redirect back + flash"),
        (520, 160, 940, "Оновлена сторінка"),
    ]
    for x1, x2, y, text in steps:
        arrow(draw, (x1, y), (x2, y))
        draw.text((min(x1, x2) + 20, y - 40), text, fill="black", font=font(20))

    img.save(out)
    return out


def generate_er_diagram():
    img, draw, out = make_canvas("er_diagram.png", 1800, 1200)
    box(draw, (70, 90, 520, 410), "users\nPK id\nname\nemail\npassword\nrole\ndeleted_at", title=True)
    box(draw, (650, 90, 1080, 370), "coaches\nPK id\nFK user_id\nbio\nphone\nspecialization", title=True)
    box(draw, (1240, 90, 1710, 410), "sports\nPK id\nname\ndescription\nlocation\nFK coach_id", title=True)
    box(draw, (310, 700, 800, 1080), "registrations\nPK id\nFK user_id\nFK training_id\nstatus\ncreated_at", title=True)
    box(draw, (1120, 700, 1710, 1080), "trainings\nPK id\nFK sport_id\ndate\ntime\nplace\nnotes\nis_cancelled\nis_completed", title=True)

    arrow(draw, (520, 220), (650, 220), "1:1")
    arrow(draw, (1080, 220), (1240, 220), "1:N")
    arrow(draw, (530, 380), (420, 700), "1:N")
    arrow(draw, (1450, 410), (1450, 700), "1:N")
    arrow(draw, (800, 900), (1120, 900), "N:1")
    img.save(out)
    return out


def load_routes():
    with open(TMP / "routes.json", "r", encoding="utf-8") as f:
        return json.load(f)


def load_schema():
    schema = defaultdict(list)
    raw = (TMP / "schema.tsv").read_text(encoding="utf-8").splitlines()
    for line in raw:
        table, column, col_type, nullable, key = line.split("\t")
        schema[table].append(
            {
                "column": column,
                "type": col_type,
                "nullable": nullable,
                "key": key,
            }
        )
    return schema


def route_groups(routes):
    groups = defaultdict(int)
    for route in routes:
        uri = route["uri"]
        if uri.startswith("admin"):
            groups["Адміністративні"] += 1
        elif uri.startswith("coach"):
            groups["Кабінет тренера"] += 1
        elif uri in {"login", "register", "logout"}:
            groups["Автентифікація"] += 1
        elif uri.startswith("sports"):
            groups["Спортивні секції"] += 1
        elif uri.startswith("trainings"):
            groups["Тренування та записи"] += 1
        elif uri.startswith("profile") or uri.startswith("registrations") or uri.startswith("users"):
            groups["Профіль та учасники"] += 1
        else:
            groups["Службові та головна сторінка"] += 1
    return groups


def add_simple_table(doc, caption, headers, rows, col_widths=None):
    add_table_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph()


def add_title_page(doc):
    add_paragraph(doc, "ЗАТВЕРДЖЕНО", bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, indent=False)
    add_paragraph(doc, "Наказ ректора ДВНЗ «ПДТУ»", align=WD_ALIGN_PARAGRAPH.LEFT, indent=False)
    add_paragraph(doc, "від 30 серпня.2019 № 147-05", align=WD_ALIGN_PARAGRAPH.LEFT, indent=False)
    add_paragraph(doc, "Форма № ПДТУ-10.03", align=WD_ALIGN_PARAGRAPH.LEFT, indent=False)
    doc.add_paragraph()
    add_paragraph(doc, "Державний вищий навчальний заклад", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    add_paragraph(doc, "«Приазовський державний технічний університет»", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    doc.add_paragraph()
    add_paragraph(doc, "Кафедра «Комп’ютерні науки»", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    doc.add_paragraph()
    doc.add_paragraph()
    add_paragraph(doc, "ЗВІТ", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, size=16)
    add_paragraph(doc, "ПРО ПЕРЕДДИПЛОМНУ ПРАКТИКУ", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, size=16)
    doc.add_paragraph()
    add_paragraph(doc, "База практики: приватний фітнес-клуб, у межах якого розроблено веборієнтовану інформаційну систему організації занять спортивних секцій", indent=False)
    add_paragraph(doc, "Місцезнаходження: м. Дніпро", indent=False)
    add_paragraph(doc, "Строки практики: лютий – березень 2026 року", indent=False)
    add_paragraph(doc, "Студент: Гогоша І. О.", indent=False)
    add_paragraph(doc, "Спеціальність: 122 Комп’ютерні науки", indent=False)
    add_paragraph(doc, "Освітня програма: Комп’ютерні науки", indent=False)
    add_paragraph(doc, "Керівник практики: викладач кафедри «Комп’ютерні науки»", indent=False)
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    add_paragraph(doc, "м. Дніпро – 2026 рік", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    doc.add_page_break()


def add_contents(doc):
    add_heading(doc, "ВСТУП")
    items = [
        "1. ОПИС ТА АНАЛІЗ ОБ’ЄКТА ДОСЛІДЖЕННЯ",
        "2. ОГЛЯД АНАЛОГІВ ТА АНАЛІЗ ІСНУЮЧИХ РІШЕНЬ",
        "2.1 ОГЛЯД АНАЛОГІВ",
        "2.2 АНАЛІЗ СУЧАСНИХ ТЕХНОЛОГІЙ",
        "3. ТЕХНІЧНЕ ЗАВДАННЯ",
        "3.1 ВИМОГИ ДО СТРУКТУРИ ТА ФУНКЦІОНУВАННЯ СИСТЕМИ",
        "3.2 ВИМОГИ ДО КОРИСТУВАЧІВ СИСТЕМИ",
        "3.3 ВИМОГИ ДО ІНФОРМАЦІЙНОГО ЗАБЕЗПЕЧЕННЯ",
        "3.4 ВИМОГИ ДО ПРОГРАМНОГО ЗАБЕЗПЕЧЕННЯ",
        "3.5 ВИМОГИ ДО ТЕХНІЧНОГО ЗАБЕЗПЕЧЕННЯ",
        "4. ЕТАПИ ВИРІШЕННЯ ЗАВДАННЯ",
        "4.1 АНАЛІЗ ВИМОГ",
        "4.2 ПРОЄКТУВАННЯ СИСТЕМИ",
        "4.3 РОЗРОБКА СИСТЕМИ",
        "4.4 ТЕСТУВАННЯ ТА ВІДЛАГОДЖЕННЯ",
        "ПЕРЕЛІК ДЖЕРЕЛ",
        "ДОДАТКИ",
    ]
    for item in items:
        add_paragraph(doc, item, indent=False)
    doc.add_page_break()


def text_block(doc, paragraphs):
    for paragraph in paragraphs:
        add_paragraph(doc, paragraph)


def build_report():
    routes = load_routes()
    schema = load_schema()
    groups = route_groups(routes)

    diagram_paths = [
        generate_use_case(),
        generate_class_diagram(),
        generate_component_diagram(),
        generate_sequence_diagram(),
        generate_er_diagram(),
    ]

    doc = Document()
    configure_document(doc)
    add_title_page(doc)
    add_contents(doc)

    add_heading(doc, "ВСТУП")
    text_block(doc, [
        "Переддипломна практика була присвячена дослідженню, аналізу та узагальненню результатів розроблення веборієнтованої інформаційної системи організації занять спортивних секцій фітнес-клубу. Об’єктом дослідження виступає процес планування спортивних секцій, публікації розкладу тренувань, реєстрації користувачів на заняття, а також адміністрування довідників і обліку статусів відвідування.",
        "Актуальність теми зумовлена потребою фітнес-клубів і спортивних організацій швидко обслуговувати велику кількість клієнтів, надавати їм зручний доступ до розкладу занять і забезпечувати прозорий механізм взаємодії між адміністраторами, тренерами та відвідувачами. Традиційний облік у вигляді телефонних записів, електронних таблиць або повідомлень у месенджерах ускладнює підтримання актуальності інформації та не дає змоги оперативно формувати статистику.",
        "Метою практики є дослідження структури та функціональних можливостей програмного продукту, побудованого на основі Laravel, Inertia.js та Vue 3, визначення його архітектурних особливостей, аналіз використаних технологій, формування технічного завдання та опис етапів реалізації системи. Практичним результатом є підготовка цілісного опису рішення на основі реального репозиторію з маршрутами, контролерами, моделями, міграціями, клієнтськими сторінками та допоміжною інфраструктурою Docker.",
        "У ході виконання практики було проаналізовано структуру репозиторію, серверну та клієнтську частини застосунку, навігацію, базу даних, ролі користувачів, модулі адміністрування, механізми локалізації, інструменти контейнеризації та сценарії заповнення тестовими даними. Окрему увагу приділено побудові діаграм, які відображають реальні сутності предметної області та взаємодію програмних компонентів.",
    ])

    add_heading(doc, "1. ОПИС ТА АНАЛІЗ ОБ’ЄКТА ДОСЛІДЖЕННЯ")
    text_block(doc, [
        "Предметна область системи охоплює діяльність фітнес-клубу, у якому існують спортивні секції, окремі тренування, тренери та користувачі, що записуються на заняття. Центральною задачею є автоматизація процесу організації занять: від створення секції та призначення тренера до реєстрації учасників, контролю статусів запису та фіксації завершення тренування.",
        "У межах дослідженого проєкту реалізовано вебзастосунок із розподілом прав доступу за ролями. Користувач без спеціальних повноважень може реєструватися, входити до системи, переглядати перелік секцій, працювати з календарем тренувань, переглядати деталі окремих занять і керувати власними записами у профілі. Тренер отримує персональний кабінет із переліком закріплених тренувань та реєстрацій. Адміністратор має доступ до панелі керування користувачами, тренерами, секціями, тренуваннями й записами.",
        "Серверна частина побудована на Laravel 10 і використовує класичну MVC-організацію коду. Маршрути визначено в `routes/web.php` і згруповано за публічним, адміністративним та тренерським сегментами. Для обмеження доступу застосовано middleware `Authenticate`, `AdminMiddleware`, `CoachMiddleware` та `SetLocale`. На рівні моделей використовується Eloquent ORM, що описує основні сутності: `User`, `Coach`, `Sport`, `Training`, `Registration`.",
        "Клієнтський інтерфейс реалізовано на Vue 3 у зв’язці з Inertia.js, що дає змогу будувати SPA-подібний досвід без окремого REST API для більшості сторінок. Навігація між сторінками здійснюється через Inertia, а рендеринг виконується всередині Blade-шаблону `resources/views/app.blade.php`. Це дозволяє поєднати маршрути Laravel з реактивним інтерфейсом Vue та зберегти єдину точку входу.",
        "Із точки зору функціональних модулів система містить блок автентифікації, публічний каталог спортивних секцій, календар тренувань, сторінки перегляду деталей, модуль реєстрацій, кабінет користувача, адміністративну панель і кабінет тренера. Важливою властивістю рішення є підтримка кількох ролей у межах єдиного коду та єдиної бази даних, що спрощує супровід і розширення системи.",
    ])

    add_simple_table(
        doc,
        "Таблиця 1.1 – Основні функціональні модулі системи",
        ["Модуль", "Призначення", "Ключові сторінки"],
        [
            ["Автентифікація", "Реєстрація, вхід, завершення сеансу", "/register, /login"],
            ["Спортивні секції", "Пошук, перегляд та фільтрація секцій", "/sports, /sports/{id}"],
            ["Тренування", "Календар занять і детальна сторінка тренування", "/trainings, /trainings/{id}"],
            ["Реєстрації", "Запис, повторний запис, скасування", "POST /trainings/*, /profile"],
            ["Адміністрування", "CRUD користувачів, тренерів, секцій, тренувань, записів", "/admin/*"],
            ["Кабінет тренера", "Огляд власних тренувань та статусів учасників", "/coach, /coach/trainings, /coach/registrations"],
        ],
    )

    add_simple_table(
        doc,
        "Таблиця 1.2 – Розподіл маршрутів за групами",
        ["Група", "Кількість маршрутів"],
        [[group, count] for group, count in groups.items()],
    )

    for caption, path in [
        ("Рисунок 1.1 – Головна сторінка системи", SCREENSHOTS / "01_home.png"),
        ("Рисунок 1.2 – Сторінка входу до системи", SCREENSHOTS / "02_login.png"),
        ("Рисунок 1.3 – Каталог спортивних секцій", SCREENSHOTS / "03_sports.png"),
    ]:
        add_image(doc, path)
        add_figure_caption(doc, caption)

    text_block(doc, [
        "Маршрути публічної частини охоплюють перегляд головної сторінки, реєстрацію та авторизацію, роботу з каталогом секцій і календарем тренувань. Захищені маршрути доповнюють цей перелік персональним профілем користувача, операціями керування записами та переглядом списку учасників. Така структура є логічною для предметної області, оскільки відокремлює операції перегляду від змін даних.",
        "На головній сторінці для авторизованого користувача система відображає найближчу активну реєстрацію, а для неавторизованого відвідувача – останні додані секції. Це свідчить про використання персоналізації на базовому рівні та демонструє орієнтацію застосунку на швидкий доступ до найбільш актуальної інформації.",
    ])

    add_heading(doc, "2. ОГЛЯД АНАЛОГІВ ТА АНАЛІЗ ІСНУЮЧИХ РІШЕНЬ")
    text_block(doc, [
        "При розробленні інформаційних систем для спортивних секцій зазвичай розглядаються три типи рішень: універсальні CRM-платформи для клубів, окремі системи розкладу й бронювання та спеціалізовані вебзастосунки, побудовані під конкретний бізнес-процес. Кожен підхід має власні переваги, але не всі з них однаково добре адаптуються до невеликих фітнес-клубів, де важлива гнучкість і можливість швидко розширювати функціональність.",
    ])

    add_heading(doc, "2.1 ОГЛЯД АНАЛОГІВ")
    text_block(doc, [
        "Першу групу аналогів становлять хмарні рішення для керування клубами, які поєднують CRM, розклад, облік абонементів та маркетингові інструменти. Їх перевагою є швидкий старт і наявність готових інтеграцій, але недоліком виступають висока залежність від зовнішнього постачальника, обмеження у зміні логіки та додаткові витрати на ліцензування.",
        "Другу групу становлять прості системи онлайн-бронювання або календарні платформи, які дозволяють створювати заняття та резервувати місця. Їх основна проблема полягає у тому, що вони слабо враховують ролеву модель спортивного клубу, де адміністратор і тренер мають різні сценарії роботи, а статуси реєстрації потребують детального обліку.",
        "Третю групу утворюють індивідуальні вебзастосунки, орієнтовані на конкретну організацію. Саме до цього класу належить досліджуваний проєкт. Він дозволяє врахувати специфіку фітнес-клубу, зберегти локальний контроль над кодом і базою даних, а також змінювати модель даних без прив’язки до сторонніх сервісів.",
    ])

    add_simple_table(
        doc,
        "Таблиця 2.1 – Порівняння підходів до автоматизації спортивних секцій",
        ["Підхід", "Переваги", "Недоліки"],
        [
            ["Хмарна CRM", "Швидке впровадження, готові інтеграції", "Менша гнучкість, абонплата"],
            ["Система бронювання", "Простота та швидкий старт", "Немає повної ролевої моделі"],
            ["Індивідуальний вебзастосунок", "Адаптація до бізнес-процесу, контроль над даними", "Потрібен супровід і розвиток коду"],
        ],
    )

    add_heading(doc, "2.2 АНАЛІЗ СУЧАСНИХ ТЕХНОЛОГІЙ")
    text_block(doc, [
        "У досліджуваному рішенні використано сучасний стек веброзробки. Laravel забезпечує стабільний серверний каркас, ORM Eloquent, маршрутизацію, middleware, валідацію та інтеграцію з інфраструктурними сервісами. Vue 3 відповідає за інтерактивний інтерфейс, а Inertia.js усуває потребу у створенні окремого REST-API для більшості сторінок, зберігаючи при цьому плавну навігацію між станами.",
        "Для збирання фронтенд-ресурсів застосовується Vite, що забезпечує швидку локальну розробку та гаряче оновлення ресурсів. У структурі проєкту присутній окремий контейнер `vite`, а також налаштовано сценарії `dev` і `build` у `package.json`. Це свідчить про орієнтацію рішення на сучасний процес фронтенд-збирання.",
        "База даних MySQL використовується як основне сховище предметних сутностей. Redis залучено для кешування, сесій та черг. Для контейнеризації та відтворюваності середовища застосовується Docker Compose з контейнерами `app`, `nginx`, `db`, `redis`, `vite`, `queue`. Така інфраструктура спрощує розгортання та тестування застосунку в різних середовищах.",
        "Важливою особливістю є багатомовність інтерфейсу. У проєкті наявні окремі словники `uk`, `en` і тимчасово прихований `ru`, а на сервері працює middleware `SetLocale`, що дозволяє керувати мовою через cookie або заголовок. Для відображення дат та часу реалізовано власні форматери, пов’язані з поточною мовою інтерфейсу.",
    ])

    for idx, path in enumerate(diagram_paths[:3], start=1):
        add_image(doc, path)
        add_figure_caption(doc, f"Рисунок 2.{idx} – {'Діаграма варіантів використання' if idx == 1 else 'Діаграма класів' if idx == 2 else 'Компонентна діаграма системи'}")

    add_heading(doc, "3. ТЕХНІЧНЕ ЗАВДАННЯ")
    text_block(doc, [
        "На основі аналізу репозиторію та логіки застосунку можна сформулювати технічне завдання на створення веборієнтованої інформаційної системи організації занять спортивних секцій фітнес-клубу. Система повинна забезпечувати централізоване ведення довідників секцій, тренувань, користувачів, тренерів та реєстрацій, а також надавати рольові інтерфейси для різних категорій користувачів.",
    ])

    add_heading(doc, "3.1 ВИМОГИ ДО СТРУКТУРИ ТА ФУНКЦІОНУВАННЯ СИСТЕМИ")
    add_bullets(doc, [
        "система повинна підтримувати публічну частину, адміністративну панель і кабінет тренера;",
        "для зберігання даних повинна використовуватися реляційна база MySQL;",
        "переходи між сторінками мають виконуватися через Inertia.js без повного перезавантаження застосунку;",
        "для запису на тренування необхідно забезпечити перевірку статусів, недопущення дублювання та повторний запис після скасування;",
        "при скасуванні або завершенні тренування система повинна масово оновлювати пов’язані реєстрації;",
        "для кожної ролі мають бути реалізовані окремі маршрути та засоби контролю доступу.",
    ])

    add_heading(doc, "3.2 ВИМОГИ ДО КОРИСТУВАЧІВ СИСТЕМИ")
    text_block(doc, [
        "Користувач ролі `user` повинен мати змогу пройти реєстрацію, авторизуватися, переглядати секції та календар тренувань, записуватися на доступні заняття, повторно активувати скасовані записи та контролювати власну історію участі. Інтерфейс для цієї ролі має бути простим і зосередженим на кінцевій дії – виборі заняття.",
        "Користувач ролі `coach` повинен отримувати доступ лише до власних секцій і тренувань. Для нього передбачено перегляд статистики, редагування тренувань, скасування або завершення заняття та оновлення статусів реєстрацій учасників. Це зменшує навантаження на адміністратора та передає частину операцій безпосередньо відповідальному тренеру.",
        "Користувач ролі `admin` повинен мати повний доступ до довідників і записів. Саме адміністратор створює та редагує користувачів, тренерів, секції, тренування, а також переглядає загальну статистику системи. Доступ до цих функцій обмежується middleware `AdminMiddleware`.",
    ])

    add_heading(doc, "3.3 ВИМОГИ ДО ІНФОРМАЦІЙНОГО ЗАБЕЗПЕЧЕННЯ")
    text_block(doc, [
        "Інформаційне забезпечення системи базується на таблицях `users`, `coaches`, `sports`, `trainings`, `registrations`, а також службових таблицях Laravel. Сутності пов’язані зовнішніми ключами та формують нормалізовану структуру даних, достатню для відображення ролей, переліку секцій, календаря тренувань і статусів участі.",
        "Особливе значення має таблиця `registrations`, у якій зберігається зв’язок між користувачем і тренуванням. Вона містить перелік бізнес-статусів: `pending`, `approved`, `cancelled`, `rejected`, `attended`, `no_show`. Завдяки цьому система відстежує повний життєвий цикл реєстрації, а не лише сам факт запису.",
    ])

    add_simple_table(
        doc,
        "Таблиця 3.1 – Основні таблиці предметної області",
        ["Таблиця", "Призначення", "Ключові поля"],
        [
            ["users", "облік усіх користувачів системи", "id, name, email, role"],
            ["coaches", "дані тренерів", "user_id, specialization, phone"],
            ["sports", "довідник спортивних секцій", "name, location, coach_id"],
            ["trainings", "розклад занять", "sport_id, date, time, is_cancelled, is_completed"],
            ["registrations", "запис користувача на тренування", "user_id, training_id, status"],
        ],
    )

    add_heading(doc, "3.4 ВИМОГИ ДО ПРОГРАМНОГО ЗАБЕЗПЕЧЕННЯ")
    add_bullets(doc, [
        "серверна платформа – PHP 8.1 та Laravel 10;",
        "клієнтська частина – Vue 3, Inertia.js, Ziggy, Vite;",
        "допоміжні сервіси – Redis і MySQL 8;",
        "контейнеризація та запуск локального середовища – Docker Compose;",
        "формування ресурсів інтерфейсу – Node.js та npm.",
    ])

    add_heading(doc, "3.5 ВИМОГИ ДО ТЕХНІЧНОГО ЗАБЕЗПЕЧЕННЯ")
    text_block(doc, [
        "Для розгортання системи достатньо робочої станції або сервера, на якому доступні Docker, Docker Compose та сучасний веббраузер. У проєкті передбачені окремі контейнери для вебзастосунку, вебсервера Nginx, MySQL, Redis, Vite та черги. Це дозволяє запускати систему як у локальному середовищі розробки, так і на продуктивній інфраструктурі з мінімальними змінами конфігурації.",
    ])

    add_heading(doc, "4. ЕТАПИ ВИРІШЕННЯ ЗАВДАННЯ")
    add_heading(doc, "4.1 АНАЛІЗ ВИМОГ")
    text_block(doc, [
        "На першому етапі було визначено основних учасників процесу: відвідувач клубу, тренер і адміністратор. Для кожної ролі сформовано перелік дій, які знаходять відображення у відповідних маршрутах і сторінках. Також було встановлено, що ключовим бізнес-об’єктом є тренування, до якого прив’язані секція, тренер і множина реєстрацій користувачів.",
        "Аналіз коду контролерів показав, що значну роль відіграє узгодження статусів. Наприклад, при завершенні тренування у `CoachTrainingController` записи зі статусом `pending` або `approved` переводяться у `attended`, а записи зі статусом `cancelled` чи `rejected` – у `no_show`. Таким чином, бізнес-правила закладені безпосередньо в серверну логіку.",
    ])

    add_heading(doc, "4.2 ПРОЄКТУВАННЯ СИСТЕМИ")
    text_block(doc, [
        "Проєктування системи виконано на основі реляційної моделі даних та рольового підходу до доступу. Модель `User` є базовою для всіх категорій користувачів. Тренерська роль деталізується через окрему таблицю `coaches`, яка пов’язується із таблицею `sports`. Тренування належать секціям, а реєстрації зв’язують користувачів із тренуваннями. Така структура дозволяє реалізувати як публічні сценарії перегляду, так і службові сценарії керування.",
        "На клієнтському рівні сформовано окремі layout-компоненти `AppLayout` та `AdminLayout`. Перший використовується для загальнодоступних сторінок і містить шапку та підвал, другий – для адміністративного розділу з бічною навігацією та верхньою панеллю. Поділ layout-компонентів спрощує масштабування та підтримує візуальне розмежування між типами інтерфейсів.",
    ])
    add_image(doc, diagram_paths[3])
    add_figure_caption(doc, "Рисунок 4.1 – Діаграма послідовності запису на тренування")
    add_image(doc, diagram_paths[4])
    add_figure_caption(doc, "Рисунок 4.2 – ER-діаграма бази даних")

    add_heading(doc, "4.3 РОЗРОБКА СИСТЕМИ")
    text_block(doc, [
        "Розробка серверної частини охопила створення міграцій, моделей, контролерів і middleware. Міграції формують фізичну структуру таблиць і поступово розширюють її: від початкових `users`, `sports`, `trainings` та `registrations` до додавання `coaches`, ознак скасування та завершення тренувань, а також soft delete для користувачів. Це демонструє еволюційний підхід до розвитку схеми даних.",
        "Розробка клієнтської частини включала створення сторінок `Home`, `Sports/Index`, `Sports/Show`, `Trainings/Index`, `Trainings/Show`, `Profile/Index`, а також багатьох сторінок адміністративного й тренерського сегментів. Для переліків застосовано пагінацію, фільтрацію та реактивний пошук. У календарі тренувань реалізовано візуальний розподіл занять по днях місяця.",
        "Важливою складовою є інфраструктура контейнеризації. У `docker-compose.yml` описано сервіси `app`, `vite`, `db`, `redis`, `nginx`, `queue`, що дозволяє запускати застосунок у стандартизованому середовищі. Наявність `Dockerfile`, `Dockerfile.dev`, `QUICKSTART.md` та `DEPLOYMENT.md` свідчить, що проєкт орієнтований не лише на локальну розробку, а й на подальше розгортання.",
    ])

    for caption, path in [
        ("Рисунок 4.3 – Сторінка перегляду секції", SCREENSHOTS / "04_sport_show.png"),
        ("Рисунок 4.4 – Календар тренувань", SCREENSHOTS / "05_trainings.png"),
        ("Рисунок 4.5 – Детальна сторінка тренування", SCREENSHOTS / "06_training_show.png"),
        ("Рисунок 4.6 – Панель адміністратора", SCREENSHOTS / "07_admin_dashboard.png"),
        ("Рисунок 4.7 – Керування секціями в адміністративному розділі", SCREENSHOTS / "08_admin_sports.png"),
        ("Рисунок 4.8 – Кабінет тренера", SCREENSHOTS / "09_coach_dashboard.png"),
        ("Рисунок 4.9 – Перелік тренувань тренера", SCREENSHOTS / "10_coach_trainings.png"),
        ("Рисунок 4.10 – Профіль користувача з історією записів", SCREENSHOTS / "11_profile.png"),
    ]:
        add_image(doc, path)
        add_figure_caption(doc, caption)

    add_heading(doc, "4.4 ТЕСТУВАННЯ ТА ВІДЛАГОДЖЕННЯ")
    text_block(doc, [
        "Тестування системи виконувалося шляхом відтворення основних сценаріїв у запущеному Docker-оточенні. Було перевірено коректність відкриття публічних сторінок, вхід під ролями адміністратора, тренера та звичайного користувача, відображення статистики, переходи між сторінками та наявність seeded-даних для демонстрації роботи системи.",
        "Окрему увагу приділено перевірці ролей та middleware. Публічні сторінки відкриваються без авторизації, сторінки `/admin/*` вимагають роль адміністратора, а сторінки `/coach/*` – роль тренера. Це забезпечує базовий рівень інформаційної безпеки та розмежування повноважень.",
        "У межах відлагодження важливою є перевірка бізнес-переходів між статусами реєстрацій, оскільки саме вони впливають на достовірність статистики та історії користувача. Відлагодження також включає перевірку міграцій, сидування бази, правильності локалізації інтерфейсу та коректності формування календаря тренувань.",
    ])

    add_heading(doc, "ПЕРЕЛІК ДЖЕРЕЛ")
    sources = [
        "1. Laravel Documentation. URL: https://laravel.com/docs/10.x.",
        "2. Inertia.js Documentation. URL: https://inertiajs.com/.",
        "3. Vue.js Documentation. URL: https://vuejs.org/guide/.",
        "4. Vite Documentation. URL: https://vitejs.dev/guide/.",
        "5. Docker Documentation. URL: https://docs.docker.com/.",
        "6. MySQL 8 Reference Manual. URL: https://dev.mysql.com/doc/.",
        "7. Вихідний код досліджуваного проєкту «sport-training-app».",
    ]
    for source in sources:
        add_paragraph(doc, source, indent=False)

    doc.add_page_break()
    add_heading(doc, "ДОДАТКИ")

    add_heading(doc, "ДОДАТОК А – СТРУКТУРА МАРШРУТІВ СИСТЕМИ")
    route_rows = []
    for route in routes:
        route_rows.append([
            route["method"],
            route["uri"],
            route.get("name") or "—",
            route["action"].split("@")[0].split("\\")[-1] if "@" in route["action"] else route["action"],
        ])
    add_simple_table(
        doc,
        "Таблиця А.1 – Перелік маршрутів системи",
        ["Метод", "URI", "Назва", "Обробник"],
        route_rows,
    )

    add_heading(doc, "ДОДАТОК Б – СТРУКТУРА БАЗИ ДАНИХ СИСТЕМИ")
    schema_rows = []
    for table_name, columns in schema.items():
        for column in columns:
            schema_rows.append([
                table_name,
                column["column"],
                column["type"],
                column["nullable"],
                column["key"] or "—",
            ])
    add_simple_table(
        doc,
        "Таблиця Б.1 – Структура таблиць бази даних",
        ["Таблиця", "Поле", "Тип", "NULL", "Ключ"],
        schema_rows,
    )

    add_heading(doc, "ДОДАТОК В – ІНТЕРФЕЙС СТОРІНКИ")
    add_image(doc, SCREENSHOTS / "07_admin_dashboard.png")
    add_figure_caption(doc, "Рисунок В.1 – Інтерфейс адміністративної панелі")

    add_heading(doc, "ДОДАТОК Г – ІНТЕРФЕЙС СТОРІНКИ")
    add_image(doc, SCREENSHOTS / "09_coach_dashboard.png")
    add_figure_caption(doc, "Рисунок Г.1 – Інтерфейс кабінету тренера")

    doc.save(OUT)


if __name__ == "__main__":
    build_report()
